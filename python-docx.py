from docx import Document
from docx.shared import Inches
from PIL import Image
import requests
from docx.shared import Pt
from flask import Flask, url_for

document = Document()

document.add_picture('propertyPicsTitle.png', width=Inches(6), height=Inches(1.5))

pictures = ['property0.jpeg', 'property1.jpg', 'property2.jpg','property3.jpg', 'property4.jpg', 'property5.jpg','property2.jpg', 'property4.jpg', 'property4.jpg', 'property4.jpg', 'property4.jpg', 'property4.jpg', 'property4.jpg', 'property4.jpg','property0.jpeg', 'property1.jpg', 'property2.jpg','property3.jpg', 'property4.jpg', 'property5.jpg','property2.jpg', 'property4.jpg', 'property4.jpg', 'property4.jpg', 'property4.jpg', 'property4.jpg', 'property4.jpg', 'property4.jpg']

tables = document.tables
table = document.add_table(rows=1, cols=3)
row_cells = table.add_row().cells
n = 0;
i = 0;

for index in range(len(pictures)):
        paragraph = row_cells[n].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(pictures[index], width=Inches(1.9), height=Inches(1))
        if n % 2 == 0 and n != 0 and i < 11:
            tables = document.tables
            table = document.add_table(rows=1, cols=3)
            row_cells = table.add_row().cells
            n = 0;
            i = i + 1
        elif i < 11:
            n = n + 1
            i = i + 1
        else:
            i = 0;
            n = 0;
            document.add_picture('propertyPicsTitle.png', width=Inches(6), height=Inches(1.5))
            tables = document.tables
            table = document.add_table(rows=1, cols=3)
            row_cells = table.add_row().cells


        

document.save('demo.docx')